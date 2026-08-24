from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "assets" / "documents" / "문동준_이력서.docx"
PHOTO = ROOT / "tmp" / "resume-profile.png"
TEAL = "2387A8"
LIGHT = "F1F2F2"
INK = "202020"
MUTED = "666666"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=160, bottom=100, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value)); node.set(qn("w:type"), "dxa")


def remove_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders"); tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}"); borders.append(element)
        element.set(qn("w:val"), "nil")


def set_repeat_table_width(table, widths):
    table.autofit = False
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol"); col.set(qn("w:w"), str(width)); grid.append(col)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW"); tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths))); tbl_w.set(qn("w:type"), "dxa")
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width)); tc_w.set(qn("w:type"), "dxa")


def font(run, size=9, bold=False, color=INK):
    run.font.name = "Malgun Gothic"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size); run.bold = bold; run.font.color.rgb = RGBColor.from_string(color)


def ptext(container, text="", size=9, bold=False, color=INK, before=0, after=2, align=None):
    p = container.add_paragraph()
    p.paragraph_format.space_before = Pt(before); p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None: p.alignment = align
    font(p.add_run(text), size, bold, color)
    return p


def add_link(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink"); hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r"); rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), TEAL); rpr.append(color)
    bold = OxmlElement("w:b"); rpr.append(bold)
    size = OxmlElement("w:sz"); size.set(qn("w:val"), "16"); rpr.append(size)
    text_node = OxmlElement("w:t"); text_node.text = text
    run.append(rpr); run.append(text_node); hyperlink.append(run); paragraph._p.append(hyperlink)


def section_title(cell, title):
    p = ptext(cell, title, 10, True, MUTED, before=5, after=2)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom"); bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "14"); bottom.set(qn("w:space"), "3"); bottom.set(qn("w:color"), TEAL)
    borders.append(bottom); p_pr.append(borders)
    return p


def bullet(cell, text, size=8.2):
    p = cell.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(.42); p.paragraph_format.first_line_indent = Cm(-.2)
    p.paragraph_format.space_after = Pt(2); p.paragraph_format.line_spacing = 1.05
    font(p.add_run(text), size, False, INK)


def project(cell, title, period, subtitle, bullets):
    table = cell.add_table(rows=1, cols=2)
    remove_table_borders(table); set_repeat_table_width(table, [4700, 1500])
    set_cell_margins(table.cell(0, 0), 0, 0, 0, 0); set_cell_margins(table.cell(0, 1), 0, 0, 0, 0)
    ptext(table.cell(0, 0), title, 10, True, INK, after=0)
    ptext(table.cell(0, 1), period, 7.5, False, MUTED, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)
    ptext(cell, subtitle, 8, True, INK, after=2)
    for item in bullets: bullet(cell, item)


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21); section.page_height = Cm(29.7)
    section.top_margin = Cm(.45); section.bottom_margin = Cm(.45)
    section.left_margin = Cm(0); section.right_margin = Cm(0)
    section.header_distance = Cm(.3); section.footer_distance = Cm(.3)

    normal = doc.styles["Normal"]
    normal.font.name = "Malgun Gothic"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(9); normal.paragraph_format.space_after = Pt(2); normal.paragraph_format.line_spacing = 1.08

    # Accent header band
    band = doc.add_table(rows=1, cols=1); band.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(band); set_repeat_table_width(band, [11906])
    set_cell_shading(band.cell(0, 0), TEAL); set_cell_margins(band.cell(0, 0), 115, 300, 115, 300)
    ptext(band.cell(0, 0), "입사지원서", 11, True, "FFFFFF", after=0)

    layout = doc.add_table(rows=1, cols=2); layout.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(layout); set_repeat_table_width(layout, [3900, 8006])
    layout.rows[0].height = Cm(27.35)
    layout.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    left, right = layout.cell(0, 0), layout.cell(0, 1)
    set_cell_shading(left, LIGHT); set_cell_margins(left, 300, 300, 260, 300)
    set_cell_margins(right, 250, 440, 220, 350)
    left.vertical_alignment = right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

    # Sidebar
    photo_p = left.add_paragraph()
    photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    photo_p.paragraph_format.space_after = Pt(9)
    photo_p.add_run().add_picture(str(PHOTO), width=Cm(4.8), height=Cm(5.6))
    ptext(left, "문동준", 20, True, INK, after=1)
    ptext(left, "게임 클라이언트 개발자", 9, True, TEAL, after=10)
    section_title(left, "CONTACT")
    ptext(left, "이메일", 7.2, True, MUTED, after=0); ptext(left, "ehdwns0126@naver.com", 7.7, False, INK, after=5)
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(2); add_link(p, "GitHub", "https://github.com/MDJ0126/")
    p = left.add_paragraph(); p.paragraph_format.space_after = Pt(8); add_link(p, "기술 블로그", "https://moondongjun.tistory.com/")

    section_title(left, "핵심 역량")
    for text in ["Unity · C#", "UniTask · Async/Await", "Addressables", "Jenkins CI/CD", "TCP 통신", "uGUI · NGUI · URP"]:
        bullet(left, text, 7.8)
    section_title(left, "학력")
    ptext(left, "김포대학교", 8.2, True, INK, after=0); ptext(left, "컴퓨터공학 · 2017.02 졸업", 7.3, False, MUTED, after=5)
    ptext(left, "공항고등학교", 8.2, True, INK, after=0); ptext(left, "2011.02 졸업", 7.3, False, MUTED, after=5)
    section_title(left, "자격증")
    for text in ["정보처리기사 · 2025.09", "정보처리산업기사 · 2015.07", "MOS 2007 Master · 2015.01"]:
        bullet(left, text, 7.4)

    # Main column
    section_title(right, "프로필")
    ptext(right, "Unity와 C# 기반 모바일 게임 클라이언트 개발자로 초기 개발부터 글로벌 출시, 라이브 운영까지 경험했습니다. 아웃게임 아키텍처, 비동기 데이터 처리, Addressables 리소스 관리와 플랫폼별 빌드·배포를 담당했습니다.", 8.4, False, INK, after=6)

    section_title(right, "경력사항")
    company = right.add_table(rows=1, cols=2); remove_table_borders(company); set_repeat_table_width(company, [4700, 1500])
    set_cell_margins(company.cell(0, 0), 0, 0, 0, 0); set_cell_margins(company.cell(0, 1), 0, 0, 0, 0)
    ptext(company.cell(0, 0), "주식회사 해긴  |  게임 클라이언트 개발", 10, True, INK, after=0)
    ptext(company.cell(0, 1), "2018.09 - 재직 중", 7.5, False, MUTED, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    project(right, "라스트 헌터 K : 서울", "2023.12 - 2026.07", "헌팅 액션 · 클라이언트 코어·시스템 개발", [
        "서버 통신과 데이터 로드를 병렬화해 초기 로딩을 약 20초에서 5초로 단축",
        "UniTask·Async/Await 기반 병렬 역직렬화와 Dictionary 전환으로 데이터 처리 효율 개선",
        "Addressables 및 AAB + PAD 적용, 크래시 프리 유저 비율 90%에서 98%로 개선",
    ])
    project(right, "데미안 전기", "2020.08 - 2023.12", "수집형 RPG · 아웃게임 프레임워크·콘텐츠 개발", [
        "UI 전환·팝업·씬 관리를 포함한 아웃게임 공통 프레임워크 설계",
        "TCP 기반 길드 대항전과 주요 콘텐츠 개발, Jenkins 빌드·배포 운영",
        "APK + OBB 구조를 AAB + PAD 기반으로 마이그레이션",
    ])
    project(right, "오버독스", "2018.09 - 2020.08", "MOBA · 클라이언트 콘텐츠 개발", [
        "Animation Event 리소스 사전 로딩으로 전투 중 프레임 스파이크 완화",
        "TCP 기반 채팅 시스템과 보상형 광고 기능 개발",
    ])

    section_title(right, "플랫폼 및 배포 경험")
    ptext(right, "Android · iOS · Google Play · App Store · Galaxy Store · 원스토어", 8.2, False, INK, after=0)

    # Quiet footer
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    font(footer.add_run("문동준 | 국문 이력서"), 7, False, MUTED)

    core = doc.core_properties
    core.title = "문동준 국문 이력서"; core.author = "문동준"; core.subject = "게임 클라이언트 개발자 이력서"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
